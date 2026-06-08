"""从 data/宋词.docx 抽取乐器、朝代、流派关系并写入 Neo4j。"""
import re
import sys
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent))
from db import get_driver

ROOT = Path(__file__).resolve().parent.parent
DOCX_PATH = ROOT / "data" / "宋词.docx"

# 背景文本关键词 → 关联节点（补充语义边，丰富图谱）
BACKGROUND_LINKS = [
    ("Period", "南宋", "RELATED_TO"),
    ("CiStyle", "清雅派", "RELATED_TO"),
    ("CiStyle", "豪放派", "RELATED_TO"),
    ("Instrument", "笙", "MENTIONED_IN"),
    ("Instrument", "琵琶", "MENTIONED_IN"),
]


def parse_tables(doc: Document) -> tuple[list, list]:
    """从 Word 表格解析词体-乐器、词人-流派-朝代。"""
    ci_type_instruments = []
    poet_style_period = []

    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if "词体" in headers and "乐器" in headers:
            idx_type = headers.index("词体")
            idx_inst = headers.index("乐器")
            idx_desc = headers.index("说明") if "说明" in headers else None
            for row in table.rows[1:]:
                cells = row.cells
                citype = cells[idx_type].text.strip()
                instrument = cells[idx_inst].text.strip()
                desc = cells[idx_desc].text.strip() if idx_desc is not None else ""
                if citype and instrument:
                    ci_type_instruments.append((citype, instrument, desc))

        if "词人" in headers and "流派" in headers and "朝代" in headers:
            idx_poet = headers.index("词人")
            idx_style = headers.index("流派")
            idx_period = headers.index("朝代")
            idx_note = headers.index("备注") if "备注" in headers else None
            for row in table.rows[1:]:
                cells = row.cells
                poet = cells[idx_poet].text.strip()
                style = cells[idx_style].text.strip()
                period = cells[idx_period].text.strip()
                note = cells[idx_note].text.strip() if idx_note is not None else ""
                if poet and style and period:
                    poet_style_period.append((poet, style, period, note))

    return ci_type_instruments, poet_style_period


def extract_background(doc: Document) -> str:
    """提取「历史背景」章节正文。"""
    lines = []
    in_section = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("三、历史背景"):
            in_section = True
            continue
        if in_section:
            if text.startswith("一、") or text.startswith("二、") or text.startswith("四、"):
                break
            if text:
                lines.append(text)
    return "\n".join(lines)


def import_supplementary(driver, docx_path: Path) -> dict:
    doc = Document(docx_path)
    ci_type_instruments, poet_style_period = parse_tables(doc)
    background = extract_background(doc)

    stats = {
        "instruments": 0,
        "ci_links": 0,
        "styles": 0,
        "periods": 0,
        "poet_links": 0,
        "context_links": 0,
    }

    with driver.session() as session:
        session.run("CREATE CONSTRAINT instrument_name IF NOT EXISTS FOR (i:Instrument) REQUIRE i.name IS UNIQUE")
        session.run("CREATE CONSTRAINT style_name IF NOT EXISTS FOR (s:CiStyle) REQUIRE s.name IS UNIQUE")
        session.run("CREATE CONSTRAINT period_name IF NOT EXISTS FOR (p:Period) REQUIRE p.name IS UNIQUE")

        if background:
            session.run(
                """
                MERGE (ctx:Context {title: '南宋声乐背景'})
                SET ctx.content = $content
                """,
                content=background,
            )

            for label, name, rel_type in BACKGROUND_LINKS:
                if name not in background and label != "Period":
                    continue
                if label == "Period" and name not in background:
                    continue
                result = session.run(
                    f"""
                    MATCH (ctx:Context {{title: '南宋声乐背景'}})
                    MATCH (n:{label} {{name: $name}})
                    MERGE (n)-[r:{rel_type}]->(ctx)
                    RETURN count(r) AS n
                    """,
                    name=name,
                ).single()
                if result and result["n"]:
                    stats["context_links"] += 1

        for citype, instrument, desc in ci_type_instruments:
            session.run(
                """
                MERGE (i:Instrument {name: $instrument})
                SET i.description = $desc
                WITH i
                MATCH (c:Cipai {type: $citype})
                MERGE (c)-[:ACCOMPANIED_BY]->(i)
                """,
                citype=citype,
                instrument=instrument,
                desc=desc,
            )
            stats["instruments"] += 1
            result = session.run(
                "MATCH (c:Cipai {type: $citype}) RETURN count(c) AS n", citype=citype
            ).single()
            stats["ci_links"] += result["n"] if result else 0

        for poet, style, period, note in poet_style_period:
            session.run("MERGE (s:CiStyle {name: $style})", style=style)
            session.run("MERGE (per:Period {name: $period})", period=period)
            stats["styles"] += 1
            stats["periods"] += 1

            result = session.run(
                """
                MATCH (p:Poet {name: $poet})
                RETURN p.name AS name
                """,
                poet=poet,
            ).single()

            if result:
                session.run(
                    """
                    MATCH (p:Poet {name: $poet})
                    SET p.note = $note
                    MERGE (s:CiStyle {name: $style})
                    MERGE (per:Period {name: $period})
                    MERGE (p)-[:BELONGS_TO]->(s)
                    MERGE (p)-[:ACTIVE_IN]->(per)
                    """,
                    poet=poet,
                    style=style,
                    period=period,
                    note=note,
                )
                stats["poet_links"] += 1
            else:
                session.run(
                    """
                    MERGE (p:Poet {name: $poet})
                    SET p.note = $note
                    MERGE (s:CiStyle {name: $style})
                    MERGE (per:Period {name: $period})
                    MERGE (p)-[:BELONGS_TO]->(s)
                    MERGE (p)-[:ACTIVE_IN]->(per)
                    """,
                    poet=poet,
                    style=style,
                    period=period,
                    note=note,
                )
                stats["poet_links"] += 1

    return stats


def main():
    if not DOCX_PATH.exists():
        raise FileNotFoundError(f"找不到 Word 文档: {DOCX_PATH}，请先运行 scripts/generate_docx.py")

    driver = get_driver()
    try:
        driver.verify_connectivity()
        stats = import_supplementary(driver, DOCX_PATH)
        print(f"Word 补充数据导入完成: {stats}")
    finally:
        driver.close()


if __name__ == "__main__":
    main()
