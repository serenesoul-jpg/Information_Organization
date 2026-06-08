"""生成结构化《宋词.docx》，供 extract_docx.py 解析。"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "data" / "宋词.docx"

# 词体—乐器映射（与 CSV 中「词体」字段对应）
CI_TYPE_INSTRUMENTS = [
    ("小令", "琵琶", "小令篇幅短促，多配琵琶弹唱，节奏明快。"),
    ("中调", "筝", "中调篇幅适中，常用筝伴奏，音律婉转。"),
    ("长调", "笙", "长调又称慢词，南宋以来多与笙箫合奏，音韵悠长。"),
]

# 词人—流派—朝代（覆盖 cipai_all.csv 中出现的词人）
POET_STYLE_PERIOD = [
    ("白居易", "婉约派", "唐五代", "唐代文人词的先驱之一。"),
    ("温庭筠", "婉约派", "唐五代", "花间词派代表，辞藻华美。"),
    ("李煜", "婉约派", "唐五代", "南唐后主，词风哀婉深挚。"),
    ("晏殊", "婉约派", "北宋", "北宋初期文坛领袖，词风清丽。"),
    ("范仲淹", "婉约派", "北宋", "兼工诗文词，格调沉雄。"),
    ("柳永", "婉约派", "北宋", "慢词开创者，市井词风流行。"),
    ("欧阳修", "婉约派", "北宋", "北宋古文运动领袖，词风疏朗。"),
    ("苏轼", "豪放派", "北宋", "豪放词派奠基人，开一代词风。"),
    ("晏几道", "婉约派", "北宋", "晏殊之子，词作情深意挚。"),
    ("秦观", "婉约派", "北宋", "婉约派代表，词风柔婉。"),
    ("周邦彦", "婉约派", "北宋", "格律精严，被称为词中老杜。"),
    ("贺铸", "婉约派", "北宋", "兼师豪放与婉约。"),
    ("李清照", "婉约派", "南宋", "南渡后词风沉郁，号为易安体。"),
    ("岳飞", "豪放派", "南宋", "抗金名将，词作慷慨激昂。"),
    ("陆游", "豪放派", "南宋", "爱国词人，词风雄浑。"),
    ("辛弃疾", "豪放派", "南宋", "豪放词派集大成者。"),
    ("姜夔", "清雅派", "南宋", "骚雅词派代表，工于律吕。"),
    ("史达祖", "清雅派", "南宋", "与姜夔并称，词风清峭。"),
    ("刘克庄", "豪放派", "南宋", "江湖诗派词人，风格多样。"),
    ("蒋捷", "婉约派", "宋末", "宋末四大词人之一，遗民情怀。"),
]


def build_docx():
    doc = Document()
    title = doc.add_heading("宋词与声乐关系文献摘要", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "本文档整理宋词词体、伴奏乐器、词人流派与朝代背景，"
        "供知识图谱结构化抽取使用。"
    )

    doc.add_heading("一、词体与伴奏乐器", level=1)
    doc.add_paragraph("下表记录不同词体长度与常用伴奏乐器之间的对应关系。")
    table1 = doc.add_table(rows=1, cols=3)
    table1.style = "Table Grid"
    hdr = table1.rows[0].cells
    hdr[0].text = "词体"
    hdr[1].text = "乐器"
    hdr[2].text = "说明"
    for citype, instrument, desc in CI_TYPE_INSTRUMENTS:
        row = table1.add_row().cells
        row[0].text = citype
        row[1].text = instrument
        row[2].text = desc

    doc.add_heading("二、词人与流派、朝代", level=1)
    doc.add_paragraph("下表记录代表词人的流派归属与主要活跃时期。")
    table2 = doc.add_table(rows=1, cols=4)
    table2.style = "Table Grid"
    hdr2 = table2.rows[0].cells
    hdr2[0].text = "词人"
    hdr2[1].text = "流派"
    hdr2[2].text = "朝代"
    hdr2[3].text = "备注"
    for poet, style, period, note in POET_STYLE_PERIOD:
        row = table2.add_row().cells
        row[0].text = poet
        row[1].text = style
        row[2].text = period
        row[3].text = note

    doc.add_heading("三、历史背景", level=1)
    doc.add_paragraph(
        "靖康之乱后，宋室南渡，词坛风气为之一变。"
        "笙与词作深度融合，促成了南宋骚雅慢词的鼎盛。"
        "姜夔、史达祖等清雅派词人精研律吕，"
        "其长调慢词多与笙箫合奏，形成独特的声乐审美。"
    )
    doc.add_paragraph(
        "北宋时期，柳永大量创制慢词，拓展了词的篇幅与表现空间；"
        "苏轼、辛弃疾则以豪放词风突破婉约传统，"
        "形成豪放与婉约两大流派并立的格局。"
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"已生成: {OUTPUT}")


if __name__ == "__main__":
    build_docx()
